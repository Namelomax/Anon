"""Read documents, anonymize them as a whole, and write the results back.

Reads ``.docx`` (paragraphs + tables + headers/footers), ``.pdf``, ``.xlsx``,
``.xml``, ``.rtf``, ``.odt`` and plain text; presentations are excluded. The
whole document is anonymized in one pass so the same entity gets the same
placeholder everywhere (e.g. one person -> ``[PERSON_1]`` across all paragraphs).
Only ``.docx`` is rebuilt structure-preserving; other formats yield anonymized
plain text.

Outputs:
* ``<name>.anon.txt``  — anonymized plain text
* ``<name>.map.json``  — placeholder -> original mapping (the deanonymize key)
* ``<name>.anon.docx`` — anonymized copy preserving paragraph/table structure
  (only for .docx inputs)
"""

from __future__ import annotations

import re
from pathlib import Path

from .engine import Anonymizer
from .mapping import Mapping, save_mapping


def read_text(path: str | Path) -> str:
    """Read a document (docx/pdf/xlsx/xml/rtf/odt/txt…) into one string."""
    path = Path(path)
    return read_text_from_bytes(path.name, path.read_bytes())


def _iter_container_paragraphs(container):
    """Yield paragraphs of a body/header/footer, including its tables."""
    for para in container.paragraphs:
        yield para
    for table in container.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    yield para


def _iter_docx_paragraphs(document):
    """Yield body paragraphs, table-cell paragraphs AND header/footer paragraphs.

    Headers/footers matter: в договорах реквизиты сторон и email часто вынесены
    в колонтитулы — раньше они не читались вовсе, поэтому не анонимизировались
    ни в тексте, ни в .docx-копии.
    """
    yield from _iter_container_paragraphs(document)
    for section in document.sections:
        for part in (
            section.header, section.footer,
            section.first_page_header, section.first_page_footer,
            section.even_page_header, section.even_page_footer,
        ):
            if part is None or getattr(part, "is_linked_to_previous", False):
                continue
            yield from _iter_container_paragraphs(part)


def _read_docx_text(path: Path) -> str:
    import docx

    document = docx.Document(str(path))
    lines = [para.text for para in _iter_docx_paragraphs(document)]
    return "\n".join(lines)


def anonymize_document(path: str | Path, anon: Anonymizer) -> tuple[str, Mapping]:
    """Read a document and anonymize its full text in one pass."""
    text = read_text(path)
    res = anon.anonymize(text)
    return res.anonymized_text, res.mapping


def _replacer(mapping: Mapping):
    """Build a function that swaps original values for their placeholders.

    Longest originals first so a value that contains a shorter one is replaced
    whole. Used to anonymize .docx paragraphs while reusing the document-wide
    mapping (consistent placeholders).
    """
    items = sorted(mapping.items(), key=lambda kv: len(kv[1]), reverse=True)
    if not items:
        return lambda s: s
    pattern = re.compile("|".join(re.escape(orig) for _, orig in items))
    inverse = {orig: ph for ph, orig in items}

    def replace(s: str) -> str:
        return pattern.sub(lambda m: inverse[m.group(0)], s)

    return replace


def _rewrite_docx(document, mapping: Mapping) -> None:
    """In-place: replace original values with placeholders in every paragraph.

    Structure is preserved; each paragraph is rewritten into a single run
    (intra-paragraph formatting is not retained — fine for a redacted artifact).
    """
    replace = _replacer(mapping)
    for para in _iter_docx_paragraphs(document):
        if not para.text:
            continue
        new_text = replace(para.text)
        if new_text == para.text:
            continue
        for run in list(para.runs):
            run.text = ""
        if para.runs:
            para.runs[0].text = new_text
        else:
            para.add_run(new_text)


def write_anonymized_docx(src: str | Path, dst: str | Path, mapping: Mapping) -> None:
    """Write an anonymized .docx copy to ``dst`` (path)."""
    import docx

    document = docx.Document(str(src))
    _rewrite_docx(document, mapping)
    document.save(str(dst))


# File formats we can extract text from. Presentations (.pptx/.ppt) are
# intentionally excluded (per requirement). Anonymization runs on the extracted
# text; the .docx path additionally rebuilds a structure-preserving copy, other
# formats are returned as anonymized .txt.
_TEXT_EXT = {".txt", ".csv", ".md", ".log", ".json"}
_UNSUPPORTED = {".pptx", ".ppt"}


def _decode(data: bytes) -> str:
    """Best-effort decode of a plain-text/byte blob (Russian docs are often cp1251)."""
    for enc in ("utf-8", "utf-8-sig", "cp1251", "latin-1"):
        try:
            return data.decode(enc)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


def _read_docx_bytes(data: bytes) -> str:
    import io

    import docx

    document = docx.Document(io.BytesIO(data))
    return "\n".join(p.text for p in _iter_docx_paragraphs(document))


def _read_pdf_bytes(data: bytes) -> str:
    import io

    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    return "\n".join((page.extract_text() or "") for page in reader.pages)


def _read_xlsx_bytes(data: bytes) -> str:
    import io

    import openpyxl

    wb = openpyxl.load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    lines: list[str] = []
    for ws in wb.worksheets:
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) for c in row if c is not None and str(c).strip()]
            if cells:
                lines.append("\t".join(cells))
    return "\n".join(lines)


def _read_xml_bytes(data: bytes) -> str:
    from lxml import etree

    root = etree.fromstring(data)  # handles the encoding declaration itself
    return "\n".join(t.strip() for t in root.itertext() if t and t.strip())


def _read_odt_bytes(data: bytes) -> str:
    import io
    import zipfile

    from lxml import etree

    with zipfile.ZipFile(io.BytesIO(data)) as z:
        content = z.read("content.xml")
    root = etree.fromstring(content)
    # One line per paragraph/heading element (tags end with "}p" / "}h").
    lines: list[str] = []
    for el in root.iter():
        tag = etree.QName(el).localname if isinstance(el.tag, str) else ""
        if tag in ("p", "h"):
            txt = "".join(el.itertext()).strip()
            if txt:
                lines.append(txt)
    return "\n".join(lines)


def _read_rtf_bytes(data: bytes) -> str:
    """Minimal RTF -> text: enough to extract content for anonymization."""
    text = data.decode("latin-1", errors="ignore")
    text = re.sub(r"\\'([0-9a-fA-F]{2})",
                  lambda m: bytes([int(m.group(1), 16)]).decode("cp1251", "ignore"), text)
    text = re.sub(r"\\u(-?\d+)\??", lambda m: chr(int(m.group(1)) % 0x10000), text)
    text = re.sub(r"\\(?:par|line|pard|sect)\b", "\n", text)
    text = re.sub(r"\\[a-zA-Z]+-?\d* ?", "", text)  # drop remaining control words
    text = text.replace("{", "").replace("}", "")
    return re.sub(r"\n{3,}", "\n\n", text).strip()


def _read_xls_bytes(data: bytes) -> str:
    """Старый Excel 97-2003 (.xls) через xlrd (openpyxl читает только .xlsx).

    xlrd 2.x поддерживает именно .xls; ставится `pip install xlrd`.
    """
    import xlrd

    book = xlrd.open_workbook(file_contents=data)
    lines: list[str] = []
    for sh in book.sheets():
        for r in range(sh.nrows):
            cells = [str(c.value).strip() for c in sh.row(r) if str(c.value).strip()]
            if cells:
                lines.append("\t".join(cells))
    return "\n".join(lines)


def _read_doc_bytes(data: bytes) -> str:
    """Старый Word 97-2003 (.doc): извлекаем текст системной утилитой.

    Чистого-python парсера для бинарного .doc нет, поэтому пробуем по очереди
    antiword → catdoc → LibreOffice (что установлено на сервере). Если ничего
    нет — понятная ошибка с подсказкой, а не бинарный мусор.
    """
    import os
    import shutil
    import subprocess
    import tempfile

    with tempfile.NamedTemporaryFile(suffix=".doc", delete=False) as tf:
        tf.write(data)
        path = tf.name
    try:
        if shutil.which("antiword"):
            r = subprocess.run(["antiword", path], capture_output=True, timeout=60)
            if r.returncode == 0 and r.stdout.strip():
                return r.stdout.decode("utf-8", "replace")
        if shutil.which("catdoc"):
            r = subprocess.run(["catdoc", "-d", "utf-8", path], capture_output=True, timeout=60)
            if r.returncode == 0 and r.stdout.strip():
                return r.stdout.decode("utf-8", "replace")
        soffice = shutil.which("soffice") or shutil.which("libreoffice")
        if soffice:
            outdir = tempfile.mkdtemp()
            subprocess.run(
                [soffice, "--headless", "--convert-to", "txt:Text", "--outdir", outdir, path],
                capture_output=True, timeout=120,
            )
            for f in os.listdir(outdir):
                if f.endswith(".txt"):
                    with open(os.path.join(outdir, f), encoding="utf-8", errors="replace") as fh:
                        return fh.read()
        raise ValueError(
            "Не удалось прочитать .doc: на сервере нет antiword / catdoc / "
            "LibreOffice. Проще всего пересохранить файл как .docx "
            "(Word/LibreOffice → «Сохранить как → .docx») и загрузить заново."
        )
    finally:
        try:
            os.unlink(path)
        except OSError:
            pass


def read_text_from_bytes(name: str, data: bytes) -> str:
    """Extract plain text from in-memory document bytes (no temp file).

    Supported: .docx .doc .pdf .xlsx/.xlsm .xls .xml .rtf .odt and plain text.
    Presentations (.pptx/.ppt) are rejected. Unknown extensions fall back to a
    best-effort text decode so nothing silently returns empty.
    """
    ext = Path(name).suffix.lower()
    if ext in _UNSUPPORTED:
        raise ValueError(f"Формат {ext} не поддерживается (презентации исключены).")
    if ext == ".docx":
        return _read_docx_bytes(data)
    if ext == ".doc":
        return _read_doc_bytes(data)
    if ext == ".pdf":
        return _read_pdf_bytes(data)
    if ext in (".xlsx", ".xlsm"):
        return _read_xlsx_bytes(data)
    if ext == ".xls":
        return _read_xls_bytes(data)
    if ext == ".xml":
        return _read_xml_bytes(data)
    if ext == ".odt":
        return _read_odt_bytes(data)
    if ext == ".rtf":
        return _read_rtf_bytes(data)
    return _decode(data)


def anonymized_docx_bytes(src_data: bytes, mapping: Mapping) -> bytes:
    """Return an anonymized .docx as bytes, built from source .docx bytes."""
    import io

    import docx

    document = docx.Document(io.BytesIO(src_data))
    _rewrite_docx(document, mapping)
    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()


def deanonymized_docx_bytes(src_data: bytes, mapping: Mapping) -> bytes:
    """Restore originals in a .docx (replace placeholders -> values), return bytes."""
    import io

    import docx

    from .deanonymize import deanonymize

    document = docx.Document(io.BytesIO(src_data))
    for para in _iter_docx_paragraphs(document):
        if not para.text:
            continue
        new_text = deanonymize(para.text, mapping)
        if new_text == para.text:
            continue
        for run in list(para.runs):
            run.text = ""
        if para.runs:
            para.runs[0].text = new_text
        else:
            para.add_run(new_text)
    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()


def anonymize_to_files(
    path: str | Path, anon: Anonymizer, out_dir: str | Path | None = None
) -> dict[str, Path]:
    """Anonymize a document and write .anon.txt, .map.json, (and .anon.docx).

    Returns a dict of the written paths plus the in-memory result counts.
    """
    path = Path(path)
    out_dir = Path(out_dir) if out_dir else path.parent
    out_dir.mkdir(parents=True, exist_ok=True)
    stem = path.stem

    anon_text, mapping = anonymize_document(path, anon)

    txt_path = out_dir / f"{stem}.anon.txt"
    map_path = out_dir / f"{stem}.map.json"
    txt_path.write_text(anon_text, encoding="utf-8")
    save_mapping(mapping, map_path)

    written = {"text": txt_path, "mapping": map_path}
    if path.suffix.lower() == ".docx":
        docx_path = out_dir / f"{stem}.anon.docx"
        write_anonymized_docx(path, docx_path, mapping)
        written["docx"] = docx_path
    return written
